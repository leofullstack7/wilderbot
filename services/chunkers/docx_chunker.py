from typing import List, Dict, Any
from docx import Document
from utils.text_utils import normalize_text, split_into_chunks

def parse_docx(content_bytes: bytes, doc_title: str) -> List[Dict[str, Any]]:
    """
    Devuelve una lista de items de texto listos para upsert:
    [{ "text": "...", "metadata": {...} }]
    - Párrafos del documento
    - Tablas: 1 fila -> 1 item (inyectando encabezados)
    """
    from io import BytesIO
    doc = Document(BytesIO(content_bytes))

    items: List[Dict[str, Any]] = []

    # 1) Texto por párrafos
    paras = []
    for p in doc.paragraphs:
        t = normalize_text(p.text)
        if t:
            paras.append(t)
    plain_text = "\n".join(paras).strip()

    for chunk in split_into_chunks(plain_text, max_chars=3000, overlap=400):
        items.append({
            "text": chunk,
            "metadata": {"doc_title": doc_title, "source_type": "word_text"}
        })

    # 2) Tablas -> cada fila un registro
    for table in doc.tables:
        headers = []
        if table.rows:
            headers = [normalize_text(c.text) for c in table.rows[0].cells]

        for r_i, row in enumerate(table.rows[1:], start=1):
            cells = [normalize_text(c.text) for c in row.cells]
            # Inyectamos encabezados como contexto
            pairs = []
            for h, v in zip(headers, cells):
                if h or v:
                    pairs.append(f"{h or 'Col'}: {v}")
            row_text = "; ".join([p for p in pairs if p])
            row_text = normalize_text(row_text)
            if row_text:
                items.append({
                    "text": row_text,
                    "metadata": {"doc_title": doc_title, "source_type": "word_table", "row_index": r_i}
                })
    return items
